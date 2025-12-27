"""
title: Document Generator
author: assistant
description: Creates downloadable Word (.docx) and Excel (.xlsx) documents from text/data
requirements: python-docx, openpyxl
version: 1.0.0
license: MIT
"""

import os
import uuid
import base64
from io import BytesIO
from typing import Callable, Awaitable, Any, Optional, List, Dict
from pydantic import BaseModel, Field
from datetime import datetime


class Tools:
    class Valves(BaseModel):
        pass

    def __init__(self):
        self.valves = self.Valves()

    async def create_word_document(
        self,
        title: str,
        content: str,
        filename: Optional[str] = None,
        __user__: dict = None,
        __event_emitter__: Callable[[dict], Awaitable[None]] = None,
    ) -> str:
        """
        Creates a downloadable Microsoft Word (.docx) document.
        
        :param title: The title/heading to display at the top of the document
        :param content: The main text content to include in the document. Use \\n for new paragraphs.
        :param filename: Optional filename (without extension). Defaults to 'document_TIMESTAMP'
        :return: A download link for the generated Word document
        """
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        if __event_emitter__:
            await __event_emitter__({
                "type": "status",
                "data": {"description": "Creating Word document...", "done": False}
            })
        
        # Create document
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content paragraphs
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
        
        # Generate filename
        if not filename:
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        filename = f"{filename}.docx"
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Encode as base64 data URI for download
        b64_content = base64.b64encode(buffer.read()).decode('utf-8')
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        data_uri = f"data:{mime_type};base64,{b64_content}"
        
        if __event_emitter__:
            await __event_emitter__({
                "type": "status",
                "data": {"description": f"Document '{filename}' created!", "done": True}
            })
        
        # Return markdown with download link
        return f"""✅ **Word document created successfully!**

📄 **Filename:** {filename}

To download, copy this data URI and paste it in your browser, or use this HTML to create a download link:

```html
<a href="{data_uri}" download="{filename}">Download {filename}</a>
```

Alternatively, here's the base64 content that can be decoded:
The document has been generated with {len(paragraphs)} paragraph(s)."""


    async def create_excel_spreadsheet(
        self,
        data: List[List[Any]],
        sheet_name: str = "Sheet1",
        filename: Optional[str] = None,
        include_header: bool = True,
        __user__: dict = None,
        __event_emitter__: Callable[[dict], Awaitable[None]] = None,
    ) -> str:
        """
        Creates a downloadable Excel (.xlsx) spreadsheet from tabular data.
        
        :param data: A list of rows, where each row is a list of cell values. First row is treated as headers if include_header is True.
        :param sheet_name: Name for the worksheet (default: 'Sheet1')
        :param filename: Optional filename (without extension). Defaults to 'spreadsheet_TIMESTAMP'
        :param include_header: If True, the first row will be formatted as a header row
        :return: A download link for the generated Excel file
        """
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        
        if __event_emitter__:
            await __event_emitter__({
                "type": "status",
                "data": {"description": "Creating Excel spreadsheet...", "done": False}
            })
        
        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name
        
        # Add data
        for row_idx, row in enumerate(data, start=1):
            for col_idx, value in enumerate(row, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                
                # Style header row
                if row_idx == 1 and include_header:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="DAEEF3", end_color="DAEEF3", fill_type="solid")
                    cell.alignment = Alignment(horizontal='center')
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Generate filename
        if not filename:
            filename = f"spreadsheet_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        filename = f"{filename}.xlsx"
        
        # Save to buffer
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        # Encode as base64 data URI
        b64_content = base64.b64encode(buffer.read()).decode('utf-8')
        mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        data_uri = f"data:{mime_type};base64,{b64_content}"
        
        if __event_emitter__:
            await __event_emitter__({
                "type": "status",
                "data": {"description": f"Spreadsheet '{filename}' created!", "done": True}
            })
        
        num_rows = len(data)
        num_cols = len(data[0]) if data else 0
        
        return f"""✅ **Excel spreadsheet created successfully!**

📊 **Filename:** {filename}
📈 **Size:** {num_rows} rows × {num_cols} columns

To download, copy this data URI and paste it in your browser:

```html
<a href="{data_uri}" download="{filename}">Download {filename}</a>
```"""
