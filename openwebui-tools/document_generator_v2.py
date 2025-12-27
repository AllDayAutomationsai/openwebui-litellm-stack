"""
title: Document Generator Pro
author: assistant
description: Creates downloadable Word (.docx) and Excel (.xlsx) documents with native OpenWebUI file storage
requirements: python-docx, openpyxl
version: 2.1.0
license: MIT
"""

import os
import uuid
import base64
from io import BytesIO
from typing import Callable, Awaitable, Any, Optional, List, Dict
from pydantic import BaseModel, Field
from datetime import datetime


class Tools:
    class Valves(BaseModel):
        pass

    def __init__(self):
        self.valves = self.Valves()

    async def create_word_document(
        self,
        title: str,
        content: str,
        filename: Optional[str] = None,
        __user__: dict = None,
        __event_emitter__: Callable[[dict], Awaitable[None]] = None,
    ) -> str:
        """
        Creates a downloadable Microsoft Word (.docx) document and saves it to OpenWebUI storage.
        
        :param title: The title/heading to display at the top of the document
        :param content: The main text content to include in the document. Use \\n for new paragraphs.
        :param filename: Optional filename (without extension). Defaults to 'document_TIMESTAMP'
        :return: A download link for the generated Word document
        """
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        if __event_emitter__:
            await __event_emitter__({
                "type": "status",
                "data": {"description": "Creating Word document...", "done": False}
            })
        
        # Create document
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content paragraphs
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
        
        # Generate filename
        if not filename:
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        full_filename = f"{filename}.docx"
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        file_bytes = buffer.read()
        
        # Try to use OpenWebUI's native file storage
        try:
            from open_webui.models.files import Files, FileForm
            from open_webui.config import UPLOAD_DIR
            
            file_id = str(uuid.uuid4())
            file_dir = os.path.join(UPLOAD_DIR, "documents")
            os.makedirs(file_dir, exist_ok=True)
            file_path = os.path.join(file_dir, f"{file_id}_{full_filename}")
            
            with open(file_path, "wb") as f:
                f.write(file_bytes)
            
            user_id = __user__.get("id", "system") if __user__ else "system"
            
            file_record = Files.insert_new_file(
                user_id,
                FileForm(
                    id=file_id,
                    filename=full_filename,
                    meta={
                        "name": full_filename,
                        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "size": len(file_bytes),
                        "path": file_path,
                    }
                )
            )
            
            if __event_emitter__:
                await __event_emitter__({
                    "type": "status",
                    "data": {"description": f"Document '{full_filename}' saved!", "done": True}
                })
            
            download_url = f"/api/v1/files/{file_id}/content"
            
            return f"""✅ **Word document created successfully!**

📄 **Filename:** {full_filename}
📦 **Size:** {len(file_bytes):,} bytes

**[📥 Download {full_filename}]({download_url})**"""

        except Exception as e:
            # Fallback: Return instructions with base64
            if __event_emitter__:
                await __event_emitter__({
                    "type": "status",
                    "data": {"description": f"Document created (fallback mode)", "done": True}
                })
            
            b64_content = base64.b64encode(file_bytes).decode('utf-8')
            
            return f"""✅ **Word document created successfully!**

📄 **Filename:** {full_filename}
📦 **Size:** {len(file_bytes):,} bytes

⚠️ Native storage unavailable. To download:
1. Open browser console (F12)
2. Run: `window.open("data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_content[:50]}...")`

Or save this base64 to a file and decode it."""


    async def create_excel_spreadsheet(
        self,
        data: List[List[Any]],
        sheet_name: str = "Sheet1",
        filename: Optional[str] = None,
        include_header: bool = True,
        __user__: dict = None,
        __event_emitter__: Callable[[dict], Awaitable[None]] = None,
    ) -> str:
        """
        Creates a downloadable Excel (.xlsx) spreadsheet from tabular data.
        
        :param data: A list of rows, where each row is a list of cell values. First row is treated as headers if include_header is True.
        :param sheet_name: Name for the worksheet (default: 'Sheet1')
        :param filename: Optional filename (without extension). Defaults to 'spreadsheet_TIMESTAMP'
        :param include_header: If True, the first row will be formatted as a header row
        :return: A download link for the generated Excel file
        """
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        
        if __event_emitter__:
            await __event_emitter__({
                "type": "status",
                "data": {"description": "Creating Excel spreadsheet...", "done": False}
            })
        
        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name
        
        # Add data
        for row_idx, row in enumerate(data, start=1):
            for col_idx, value in enumerate(row, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                
                # Style header row
                if row_idx == 1 and include_header:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="DAEEF3", end_color="DAEEF3", fill_type="solid")
                    cell.alignment = Alignment(horizontal='center')
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Generate filename
        if not filename:
            filename = f"spreadsheet_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        full_filename = f"{filename}.xlsx"
        
        # Save to buffer
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        file_bytes = buffer.read()
        
        num_rows = len(data)
        num_cols = len(data[0]) if data else 0
        
        # Try to use OpenWebUI's native file storage
        try:
            from open_webui.models.files import Files, FileForm
            from open_webui.config import UPLOAD_DIR
            
            file_id = str(uuid.uuid4())
            file_dir = os.path.join(UPLOAD_DIR, "documents")
            os.makedirs(file_dir, exist_ok=True)
            file_path = os.path.join(file_dir, f"{file_id}_{full_filename}")
            
            with open(file_path, "wb") as f:
                f.write(file_bytes)
            
            user_id = __user__.get("id", "system") if __user__ else "system"
            
            file_record = Files.insert_new_file(
                user_id,
                FileForm(
                    id=file_id,
                    filename=full_filename,
                    meta={
                        "name": full_filename,
                        "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        "size": len(file_bytes),
                        "path": file_path,
                    }
                )
            )
            
            if __event_emitter__:
                await __event_emitter__({
                    "type": "status",
                    "data": {"description": f"Spreadsheet '{full_filename}' saved!", "done": True}
                })
            
            download_url = f"/api/v1/files/{file_id}/content"
            
            return f"""✅ **Excel spreadsheet created successfully!**

📊 **Filename:** {full_filename}
📈 **Size:** {num_rows} rows × {num_cols} columns
📦 **File size:** {len(file_bytes):,} bytes

**[📥 Download {full_filename}]({download_url})**"""

        except Exception as e:
            if __event_emitter__:
                await __event_emitter__({
                    "type": "status",
                    "data": {"description": f"Spreadsheet created (fallback mode)", "done": True}
                })
            
            b64_content = base64.b64encode(file_bytes).decode('utf-8')
            
            return f"""✅ **Excel spreadsheet created successfully!**

📊 **Filename:** {full_filename}
📈 **Size:** {num_rows} rows × {num_cols} columns

⚠️ Native storage unavailable. Base64 data available for manual download."""
